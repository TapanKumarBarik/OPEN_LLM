from docx import Document
from .base import BaseDocumentParser, DocumentChunk
from typing import List, Dict

class DocxParser(BaseDocumentParser):
    def parse(self, file_path: str, metadata: Dict) -> List[DocumentChunk]:
        chunks = []
        doc = Document(file_path)
        
        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            if not text:
                continue
                
            para_metadata = metadata.copy()
            para_metadata.update({
                'paragraph_number': para_num,
                'total_paragraphs': len(doc.paragraphs)
            })
            
            para_chunks = self._create_chunks(text, para_metadata)
            chunks.extend(para_chunks)
            
        return chunks